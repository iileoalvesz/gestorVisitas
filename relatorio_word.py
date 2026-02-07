"""
Módulo para gerar relatórios em Word usando template
"""
import os
from datetime import datetime
from typing import List, Dict
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


class GeradorRelatorioWord:
    def __init__(self, template_path: str = "templates_word/RelatorioConsolidado.docx"):
        self.template_path = template_path

    def gerar_relatorio(self, visitas: List[Dict], escolas_info: Dict = None,
                       mediadores_info: Dict = None, output_path: str = None) -> str:
        """
        Gera relatório Word usando template

        Args:
            visitas: Lista de visitas para incluir no relatório
            escolas_info: Dicionário com informações das escolas
            mediadores_info: Dicionário com informações dos mediadores
            output_path: Caminho para salvar o relatório

        Returns:
            Caminho do arquivo gerado
        """
        # Carrega template
        if os.path.exists(self.template_path):
            doc = Document(self.template_path)
        else:
            doc = Document()

        # Substitui placeholders no documento
        self._substituir_placeholders(doc, visitas, escolas_info, mediadores_info)

        # Gera nome do arquivo
        if output_path is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = f"relatorios/relatorio_consolidado_{timestamp}.docx"

        os.makedirs(os.path.dirname(output_path), exist_ok=True)

        # Salva documento
        doc.save(output_path)

        return output_path

    def _substituir_placeholders(self, doc, visitas, escolas_info, mediadores_info):
        """Substitui placeholders no documento"""
        # Dicionário de substituições
        substituicoes = {
            '{{DATA_GERACAO}}': datetime.now().strftime("%d/%m/%Y"),
            '{{TOTAL_VISITAS}}': str(len(visitas)),
            '{{PERIODO}}': self._obter_periodo(visitas)
        }

        # Substitui em parágrafos
        for paragrafo in doc.paragraphs:
            for chave, valor in substituicoes.items():
                if chave in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace(chave, valor)

        # Substitui em tabelas
        for tabela in doc.tables:
            for linha in tabela.rows:
                for celula in linha.cells:
                    for chave, valor in substituicoes.items():
                        if chave in celula.text:
                            celula.text = celula.text.replace(chave, valor)

        # Adiciona tabela de visitas se houver placeholder
        self._adicionar_tabela_visitas(doc, visitas)

    def _obter_periodo(self, visitas):
        """Obtém período das visitas"""
        if not visitas:
            return "Sem visitas"

        datas = [v['data'] for v in visitas]
        data_min = min(datas)
        data_max = max(datas)

        if data_min == data_max:
            return datetime.strptime(data_min, "%Y-%m-%d").strftime("%d/%m/%Y")
        else:
            data_min_fmt = datetime.strptime(data_min, "%Y-%m-%d").strftime("%d/%m/%Y")
            data_max_fmt = datetime.strptime(data_max, "%Y-%m-%d").strftime("%d/%m/%Y")
            return f"{data_min_fmt} a {data_max_fmt}"

    def _adicionar_tabela_visitas(self, doc, visitas):
        """Adiciona tabela com detalhes das visitas"""
        # Procura por placeholder de tabela
        for paragrafo in doc.paragraphs:
            if '{{TABELA_VISITAS}}' in paragrafo.text:
                # Remove placeholder
                paragrafo.text = ''

                # Adiciona tabela
                tabela = doc.add_table(rows=1, cols=5)
                tabela.style = 'Table Grid'

                # Cabeçalho
                celulas_cabecalho = tabela.rows[0].cells
                cabecalhos = ['Data', 'Escola', 'Mediador', 'Observações', 'Anexos']

                for i, cabecalho in enumerate(cabecalhos):
                    celulas_cabecalho[i].text = cabecalho
                    # Negrito no cabeçalho
                    run = celulas_cabecalho[i].paragraphs[0].runs[0]
                    run.bold = True

                # Dados
                for visita in visitas:
                    linha = tabela.add_row().cells
                    linha[0].text = datetime.strptime(visita['data'], "%Y-%m-%d").strftime("%d/%m/%Y")
                    linha[1].text = visita['escola_nome']
                    linha[2].text = visita.get('mediador_nome', '-')
                    linha[3].text = visita.get('observacoes', '')[:100] + ('...' if len(visita.get('observacoes', '')) > 100 else '')
                    linha[4].text = str(len(visita.get('anexos', [])))

                break

    def gerar_relatorio_detalhado(self, visitas: List[Dict], escolas: List[Dict],
                                  mediadores: List[Dict] = None) -> str:
        """
        Gera relatório detalhado com todas as informações

        Args:
            visitas: Lista de visitas
            escolas: Lista de escolas
            mediadores: Lista de mediadores (opcional)

        Returns:
            Caminho do arquivo gerado
        """
        doc = Document()

        # Título
        titulo = doc.add_heading('Relatório Consolidado de Visitas às Escolas', 0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Informações gerais
        doc.add_heading('Informações Gerais', level=1)

        p = doc.add_paragraph()
        p.add_run(f'Data de Geração: ').bold = True
        p.add_run(datetime.now().strftime("%d/%m/%Y às %H:%M"))

        p = doc.add_paragraph()
        p.add_run(f'Total de Visitas: ').bold = True
        p.add_run(str(len(visitas)))

        p = doc.add_paragraph()
        p.add_run(f'Escolas do Bloco 1: ').bold = True
        p.add_run(str(len(escolas)))

        # Visitas por Escola
        doc.add_heading('Visitas por Escola', level=1)

        # Agrupa visitas por escola
        visitas_por_escola = {}
        for visita in visitas:
            escola_nome = visita['escola_nome']
            if escola_nome not in visitas_por_escola:
                visitas_por_escola[escola_nome] = []
            visitas_por_escola[escola_nome].append(visita)

        # Cria tabela
        tabela = doc.add_table(rows=1, cols=5)
        tabela.style = 'Light Grid Accent 1'

        # Cabeçalho
        cabecalhos = ['Data', 'Escola', 'Mediador', 'Observações', 'Anexos']
        celulas_cabecalho = tabela.rows[0].cells
        for i, cabecalho in enumerate(cabecalhos):
            celulas_cabecalho[i].text = cabecalho
            celulas_cabecalho[i].paragraphs[0].runs[0].bold = True

        # Dados ordenados por escola
        for escola_nome in sorted(visitas_por_escola.keys()):
            for visita in visitas_por_escola[escola_nome]:
                linha = tabela.add_row().cells
                linha[0].text = datetime.strptime(visita['data'], "%Y-%m-%d").strftime("%d/%m/%Y")
                linha[1].text = visita['escola_nome']
                linha[2].text = visita.get('mediador_nome', '-')
                obs_texto = visita.get('observacoes', '')
                linha[3].text = obs_texto[:80] + ('...' if len(obs_texto) > 80 else '')
                linha[4].text = str(len(visita.get('anexos', [])))

        # Detalhes das Visitas
        doc.add_page_break()
        doc.add_heading('Detalhes das Visitas', level=1)

        for escola_nome in sorted(visitas_por_escola.keys()):
            doc.add_heading(escola_nome, level=2)

            for i, visita in enumerate(visitas_por_escola[escola_nome], 1):
                doc.add_heading(f'Visita #{i}', level=3)

                p = doc.add_paragraph()
                p.add_run('Data: ').bold = True
                p.add_run(datetime.strptime(visita['data'], "%Y-%m-%d").strftime("%d/%m/%Y"))

                p = doc.add_paragraph()
                p.add_run('Hora: ').bold = True
                p.add_run(visita.get('hora', '-'))

                if visita.get('mediador_nome'):
                    p = doc.add_paragraph()
                    p.add_run('Mediador: ').bold = True
                    p.add_run(visita['mediador_nome'])

                if visita.get('observacoes'):
                    p = doc.add_paragraph()
                    p.add_run('Observações:').bold = True
                    doc.add_paragraph(visita['observacoes'])

                if visita.get('anexos'):
                    p = doc.add_paragraph()
                    p.add_run(f'Anexos ({len(visita["anexos"])}):').bold = True
                    for anexo in visita['anexos']:
                        doc.add_paragraph(f"• {anexo['nome_original']}", style='List Bullet')

                doc.add_paragraph()  # Espaço

        # Salva documento
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = f"relatorios/relatorio_detalhado_{timestamp}.docx"
        os.makedirs('relatorios', exist_ok=True)
        doc.save(output_path)

        return output_path
