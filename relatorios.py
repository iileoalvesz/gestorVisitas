"""
Módulo para gerar relatórios das visitas
"""
import os
from datetime import datetime
from typing import List, Dict, Optional
import pandas as pd
from tabulate import tabulate
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


class GeradorRelatorios:
    def __init__(self, pasta_relatorios: str = "relatorios"):
        self.pasta_relatorios = pasta_relatorios
        os.makedirs(pasta_relatorios, exist_ok=True)

    def _formatar_data(self, data_str: str) -> str:
        """Formata data de YYYY-MM-DD para DD/MM/YYYY"""
        try:
            data = datetime.strptime(data_str, "%Y-%m-%d")
            return data.strftime("%d/%m/%Y")
        except:
            return data_str

    def gerar_relatorio_texto(self, visitas: List[Dict],
                             escolas: Optional[List[Dict]] = None,
                             titulo: str = "Relatório de Visitas") -> str:
        """
        Gera relatório em formato texto

        Args:
            visitas: Lista de visitas
            escolas: Lista de escolas (opcional, para informações adicionais)
            titulo: Título do relatório

        Returns:
            String com relatório formatado
        """
        linhas = []
        linhas.append("=" * 80)
        linhas.append(titulo.center(80))
        linhas.append("=" * 80)
        linhas.append(f"Gerado em: {datetime.now().strftime('%d/%m/%Y às %H:%M')}")
        linhas.append(f"Total de visitas: {len(visitas)}")
        linhas.append("=" * 80)
        linhas.append("")

        if not visitas:
            linhas.append("Nenhuma visita registrada.")
            return "\n".join(linhas)

        # Agrupa visitas por escola
        visitas_por_escola = {}
        for visita in visitas:
            escola_nome = visita['escola_nome']
            if escola_nome not in visitas_por_escola:
                visitas_por_escola[escola_nome] = []
            visitas_por_escola[escola_nome].append(visita)

        # Gera seção para cada escola
        for escola_nome in sorted(visitas_por_escola.keys()):
            visitas_escola = visitas_por_escola[escola_nome]

            linhas.append(f"\n📍 {escola_nome}")
            linhas.append("-" * 80)
            linhas.append(f"Total de visitas: {len(visitas_escola)}\n")

            for i, visita in enumerate(visitas_escola, 1):
                linhas.append(f"  Visita #{i}")
                linhas.append(f"  Data: {self._formatar_data(visita['data'])} às {visita['hora']}")

                if visita.get('observacoes'):
                    linhas.append(f"  Observações: {visita['observacoes']}")

                if visita.get('anexos'):
                    linhas.append(f"  Anexos: {len(visita['anexos'])} arquivo(s)")
                    for anexo in visita['anexos']:
                        linhas.append(f"    - {anexo['nome_original']}")

                linhas.append("")

        return "\n".join(linhas)

    def gerar_relatorio_excel(self, visitas: List[Dict],
                             arquivo: Optional[str] = None) -> str:
        """
        Gera relatório em formato Excel

        Args:
            visitas: Lista de visitas
            arquivo: Nome do arquivo (opcional)

        Returns:
            Caminho do arquivo gerado
        """
        if arquivo is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            arquivo = f"relatorio_visitas_{timestamp}.xlsx"

        caminho_completo = os.path.join(self.pasta_relatorios, arquivo)

        # Prepara dados para DataFrame
        dados = []
        for visita in visitas:
            dados.append({
                'ID': visita['id'],
                'Escola': visita['escola_nome'],
                'Data': self._formatar_data(visita['data']),
                'Hora': visita['hora'],
                'Observações': visita.get('observacoes', ''),
                'Nº Anexos': len(visita.get('anexos', []))
            })

        df = pd.DataFrame(dados)

        # Cria Excel writer
        with pd.ExcelWriter(caminho_completo, engine='openpyxl') as writer:
            # Aba principal com todas as visitas
            df.to_excel(writer, sheet_name='Todas as Visitas', index=False)

            # Aba com resumo por escola
            resumo_por_escola = df.groupby('Escola').agg({
                'ID': 'count',
                'Nº Anexos': 'sum'
            }).rename(columns={'ID': 'Total de Visitas', 'Nº Anexos': 'Total de Anexos'})

            resumo_por_escola.to_excel(writer, sheet_name='Resumo por Escola')

            # Aba com resumo por data
            resumo_por_data = df.groupby('Data').agg({
                'ID': 'count',
                'Escola': lambda x: ', '.join(sorted(set(x)))
            }).rename(columns={'ID': 'Nº Visitas', 'Escola': 'Escolas Visitadas'})

            resumo_por_data.to_excel(writer, sheet_name='Resumo por Data')

        return caminho_completo

    def gerar_relatorio_resumo(self, visitas: List[Dict],
                              estatisticas: Dict) -> str:
        """
        Gera relatório resumido com estatísticas

        Args:
            visitas: Lista de visitas
            estatisticas: Dicionário com estatísticas

        Returns:
            String com relatório formatado
        """
        linhas = []
        linhas.append("=" * 80)
        linhas.append("RELATÓRIO RESUMIDO - ESTATÍSTICAS".center(80))
        linhas.append("=" * 80)
        linhas.append("")

        linhas.append(f"📊 Total de visitas realizadas: {estatisticas['total_visitas']}")
        linhas.append(f"🏫 Total de escolas visitadas: {estatisticas['total_escolas_visitadas']}")

        if estatisticas['escola_mais_visitada']:
            linhas.append(f"⭐ Escola mais visitada: {estatisticas['escola_mais_visitada']} "
                         f"({estatisticas['max_visitas_escola']} visitas)")

        linhas.append("\n" + "-" * 80)
        linhas.append("VISITAS POR ESCOLA")
        linhas.append("-" * 80)

        # Tabela de visitas por escola
        dados_escola = []
        for escola, count in sorted(estatisticas['visitas_por_escola'].items(),
                                   key=lambda x: x[1], reverse=True):
            dados_escola.append([escola, count])

        tabela = tabulate(dados_escola, headers=['Escola', 'Nº Visitas'],
                         tablefmt='grid')
        linhas.append(tabela)

        if estatisticas['visitas_por_mes']:
            linhas.append("\n" + "-" * 80)
            linhas.append("VISITAS POR MÊS")
            linhas.append("-" * 80)

            dados_mes = []
            for mes, count in sorted(estatisticas['visitas_por_mes'].items()):
                try:
                    data = datetime.strptime(mes, "%Y-%m")
                    mes_formatado = data.strftime("%B/%Y")
                except:
                    mes_formatado = mes
                dados_mes.append([mes_formatado, count])

            tabela = tabulate(dados_mes, headers=['Mês', 'Nº Visitas'],
                             tablefmt='grid')
            linhas.append(tabela)

        return "\n".join(linhas)

    def gerar_relatorio_escola(self, escola_nome: str, visitas: List[Dict]) -> str:
        """
        Gera relatório específico para uma escola

        Args:
            escola_nome: Nome da escola
            visitas: Lista de visitas da escola

        Returns:
            String com relatório formatado
        """
        linhas = []
        linhas.append("=" * 80)
        linhas.append(f"RELATÓRIO DE VISITAS - {escola_nome}".center(80))
        linhas.append("=" * 80)
        linhas.append(f"Total de visitas: {len(visitas)}")
        linhas.append("=" * 80)
        linhas.append("")

        if not visitas:
            linhas.append("Nenhuma visita registrada para esta escola.")
            return "\n".join(linhas)

        # Ordena visitas por data
        visitas_ordenadas = sorted(visitas, key=lambda x: (x['data'], x['hora']))

        for i, visita in enumerate(visitas_ordenadas, 1):
            linhas.append(f"VISITA #{i}")
            linhas.append("-" * 80)
            linhas.append(f"ID: {visita['id']}")
            linhas.append(f"Data: {self._formatar_data(visita['data'])} às {visita['hora']}")

            if visita.get('observacoes'):
                linhas.append(f"\nObservações:")
                linhas.append(visita['observacoes'])

            if visita.get('anexos'):
                linhas.append(f"\nAnexos ({len(visita['anexos'])} arquivo(s)):")
                for anexo in visita['anexos']:
                    linhas.append(f"  - {anexo['nome_original']}")
                    linhas.append(f"    Localização: {anexo['caminho']}")

            linhas.append("=" * 80)
            linhas.append("")

        return "\n".join(linhas)

    def salvar_relatorio_texto(self, conteudo: str, nome_arquivo: str) -> str:
        """
        Salva relatório em arquivo de texto

        Args:
            conteudo: Conteúdo do relatório
            nome_arquivo: Nome do arquivo

        Returns:
            Caminho do arquivo gerado
        """
        if not nome_arquivo.endswith('.txt'):
            nome_arquivo += '.txt'

        caminho_completo = os.path.join(self.pasta_relatorios, nome_arquivo)

        with open(caminho_completo, 'w', encoding='utf-8') as f:
            f.write(conteudo)

        return caminho_completo

    def listar_escolas_sem_visita(self, todas_escolas: List[Dict],
                                  visitas: List[Dict]) -> List[str]:
        """
        Lista escolas que ainda não foram visitadas

        Args:
            todas_escolas: Lista de todas as escolas
            visitas: Lista de visitas realizadas

        Returns:
            Lista de nomes de escolas sem visita
        """
        escolas_visitadas = set(v['escola_nome'] for v in visitas)
        escolas_sem_visita = []

        for escola in todas_escolas:
            if escola['nome_usual'] not in escolas_visitadas:
                escolas_sem_visita.append(escola['nome_usual'])

        return sorted(escolas_sem_visita)

    def gerar_relatorio_word(self, visitas: List[Dict],
                            arquivo: Optional[str] = None) -> str:
        """
        Gera relatório em formato Word (.docx)

        Args:
            visitas: Lista de visitas
            arquivo: Nome do arquivo (opcional)

        Returns:
            Caminho do arquivo gerado
        """
        if arquivo is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            arquivo = f"relatorio_visitas_{timestamp}.docx"

        caminho_completo = os.path.join(self.pasta_relatorios, arquivo)

        # Cria documento
        doc = Document()

        # Adiciona título
        titulo = doc.add_heading('Relatório de Visitas às Escolas', 0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Adiciona subtítulo com data e informações
        subtitulo = doc.add_paragraph()
        subtitulo.add_run(f'Gerado em: {datetime.now().strftime("%d/%m/%Y às %H:%M")}\n').bold = True
        subtitulo.add_run(f'Total de visitas: {len(visitas)}').bold = True
        subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # Espaço

        if not visitas:
            doc.add_paragraph('Nenhuma visita registrada.')
            doc.save(caminho_completo)
            return caminho_completo

        # Agrupa visitas por escola
        visitas_por_escola = {}
        for visita in visitas:
            escola_nome = visita['escola_nome']
            if escola_nome not in visitas_por_escola:
                visitas_por_escola[escola_nome] = []
            visitas_por_escola[escola_nome].append(visita)

        # Gera seção para cada escola
        for escola_nome in sorted(visitas_por_escola.keys()):
            visitas_escola = visitas_por_escola[escola_nome]

            # Título da escola
            heading_escola = doc.add_heading(f'📍 {escola_nome}', level=1)

            # Info da escola
            info = doc.add_paragraph()
            info.add_run(f'Total de visitas: {len(visitas_escola)}').italic = True

            # Lista cada visita
            for i, visita in enumerate(visitas_escola, 1):
                # Subtítulo da visita
                doc.add_heading(f'Visita #{i}', level=2)

                # Detalhes
                p_data = doc.add_paragraph()
                p_data.add_run('Data: ').bold = True
                p_data.add_run(f"{self._formatar_data(visita['data'])} às {visita['hora']}")

                # Mediador
                if visita.get('mediador_nome'):
                    p_mediador = doc.add_paragraph()
                    p_mediador.add_run('Mediador: ').bold = True
                    p_mediador.add_run(visita['mediador_nome'])

                # Observações
                if visita.get('observacoes'):
                    p_obs = doc.add_paragraph()
                    p_obs.add_run('Observações: ').bold = True
                    doc.add_paragraph(visita['observacoes'], style='Quote')

                # Anexos
                if visita.get('anexos'):
                    p_anexos = doc.add_paragraph()
                    p_anexos.add_run(f'Anexos ({len(visita["anexos"])} arquivo(s)): ').bold = True

                    for anexo in visita['anexos']:
                        doc.add_paragraph(
                            f"• {anexo['nome_original']}",
                            style='List Bullet'
                        )

                doc.add_paragraph()  # Espaço entre visitas

            # Quebra de página entre escolas
            doc.add_page_break()

        # Remove última quebra de página
        if len(doc.element.body) > 0:
            last_element = doc.element.body[-1]
            if last_element.tag.endswith('sectPr'):
                doc.element.body.remove(last_element)

        # Salva documento
        doc.save(caminho_completo)

        return caminho_completo

    def gerar_relatorio_consolidado(self, visitas: List[Dict],
                                    arquivo_template: str = None,
                                    arquivo: Optional[str] = None) -> str:
        """
        Gera relatório consolidado articulado com fotos usando template

        Args:
            visitas: Lista de visitas
            arquivo_template: Caminho do template Word (opcional)
            arquivo: Nome do arquivo de saída (opcional)

        Returns:
            Caminho do arquivo gerado
        """
        if arquivo is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            arquivo = f"relatorio_consolidado_{timestamp}.docx"

        caminho_completo = os.path.join(self.pasta_relatorios, arquivo)

        # Se houver template, usa ele; senão cria novo documento
        if arquivo_template and os.path.exists(arquivo_template):
            doc = Document(arquivo_template)
        else:
            doc = Document()

            # Se não há template, cria cabeçalho básico
            titulo = doc.add_heading('RELATÓRIO ARTICULADO PEDAGÓGICO DE ÁREA', 0)
            titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

            subtitulo = doc.add_paragraph()
            subtitulo.add_run('Convênio 24.480/2025').bold = True
            subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Adiciona data de geração
        info = doc.add_paragraph()
        info.add_run(f'Gerado em: {datetime.now().strftime("%d/%m/%Y às %H:%M")}\n').bold = True
        info.add_run(f'Total de visitas: {len(visitas)}').bold = True

        doc.add_paragraph()  # Espaço

        if not visitas:
            doc.add_paragraph('Nenhuma visita registrada.')
            doc.save(caminho_completo)
            return caminho_completo

        # Seção 1: Identificação
        doc.add_heading('1- IDENTIFICAÇÃO', level=1)

        # Tabela de articuladores (simplificada)
        mediadores_unicos = {}
        for visita in visitas:
            if visita.get('mediador_nome'):
                med_nome = visita['mediador_nome']
                escola = visita['escola_nome']
                if med_nome not in mediadores_unicos:
                    mediadores_unicos[med_nome] = escola

        if mediadores_unicos:
            doc.add_paragraph('Articuladores Pedagógicos de Área:')
            for i, (nome, escola) in enumerate(mediadores_unicos.items(), 1):
                doc.add_paragraph(f'{i}. {nome} - {escola}', style='List Bullet')

        doc.add_paragraph()

        # Seção 2: Descritivo de oficinas
        doc.add_heading('2- DESCRITIVO DE OFICINAS E ACOMPANHAMENTOS', level=1)

        # Agrupa visitas por escola
        visitas_por_escola = {}
        for visita in visitas:
            escola_nome = visita['escola_nome']
            if escola_nome not in visitas_por_escola:
                visitas_por_escola[escola_nome] = []
            visitas_por_escola[escola_nome].append(visita)

        for escola_nome in sorted(visitas_por_escola.keys()):
            visitas_escola = visitas_por_escola[escola_nome]

            doc.add_heading(f'Unidade Escolar: {escola_nome}', level=2)

            for i, visita in enumerate(visitas_escola, 1):
                p = doc.add_paragraph()
                p.add_run(f'Visita #{i} - ').bold = True
                p.add_run(f"{self._formatar_data(visita['data'])} às {visita['hora']}")

                if visita.get('mediador_nome'):
                    p_med = doc.add_paragraph()
                    p_med.add_run('Mediador: ').bold = True
                    p_med.add_run(visita['mediador_nome'])

                if visita.get('observacoes'):
                    doc.add_paragraph(visita['observacoes'])

                doc.add_paragraph()  # Espaço

        # Seção 3: Registro fotográfico
        doc.add_heading('3- REGISTRO FOTOGRÁFICO DAS ATIVIDADES DESENVOLVIDAS NAS OFICINAS', level=1)

        # Processa fotos das visitas
        for visita in visitas:
            if visita.get('anexos'):
                escola = visita['escola_nome']

                doc.add_paragraph()
                p_escola = doc.add_paragraph()
                p_escola.add_run(f'Unidade Escolar: {escola}').bold = True

                # Filtra apenas imagens
                imagens = [a for a in visita['anexos']
                          if a['nome_original'].lower().endswith(('.jpg', '.jpeg', '.png'))]

                if imagens:
                    # Adiciona até 4 imagens por página (2x2)
                    for idx, anexo in enumerate(imagens[:4]):
                        try:
                            # Monta caminho completo do anexo
                            caminho_anexo = os.path.join('static/uploads',
                                                        os.path.basename(anexo.get('caminho', '')))

                            if os.path.exists(caminho_anexo):
                                # Adiciona imagem com largura de 7cm (aproximadamente)
                                doc.add_picture(caminho_anexo, width=Inches(2.8))

                                # Adiciona legenda
                                p_legenda = doc.add_paragraph(anexo['nome_original'])
                                p_legenda.alignment = WD_ALIGN_PARAGRAPH.CENTER

                                # Quebra de linha a cada 2 imagens
                                if (idx + 1) % 2 == 0:
                                    doc.add_paragraph()
                        except Exception as e:
                            # Se não conseguir adicionar a imagem, continua
                            print(f"Erro ao adicionar imagem: {e}")
                            pass

                doc.add_paragraph()  # Espaço entre escolas

        # Salva documento consolidado
        doc.save(caminho_completo)

        return caminho_completo

    def _set_cell_shading(self, cell, color):
        """Define cor de fundo de uma celula"""
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    def _set_cell_border(self, cell):
        """Define bordas finas em uma celula"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '</w:tcBorders>'
        )
        tcPr.append(tcBorders)

    def _add_cell_text(self, cell, text, bold=False, size=8, alignment=WD_ALIGN_PARAGRAPH.LEFT):
        """Adiciona texto formatado a uma celula"""
        cell.paragraphs[0].clear()
        cell.paragraphs[0].alignment = alignment
        run = cell.paragraphs[0].add_run(text)
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.name = 'Arial'
        cell.paragraphs[0].paragraph_format.space_before = Pt(1)
        cell.paragraphs[0].paragraph_format.space_after = Pt(1)

    # Indicadores corretos da avaliacao FUNCABES
    INDICADORES_AVALIACAO = [
        ("1. ESTRUTURA PEDAGOGICA", [
            ("ind1", "Planejamento Semanal"),
            ("ind2", "Conteudo corresponde ao planejamento"),
            ("ind3", "Conteudo corresponde ao curriculo"),
        ]),
        ("2. ORGANIZACAO DA OFICINA", [
            ("ind4", "Material adequado a atividade e faixa etaria"),
            ("ind5", "Organizacao do espaco fisico"),
        ]),
        ("3. GESTAO DA OFICINA", [
            ("ind6", "Adequacao de atividades x tempo da oficina"),
            ("ind7", "Pontualidade no atendimento"),
            ("ind8", "Desenvolvimento das atividades"),
        ]),
        ("4. CLIMA DA OFICINA", [
            ("ind9", "Estimulo as atitudes/valores"),
            ("ind10", "Relacionamento com os estudantes"),
        ]),
    ]

    ESCALA_VALORES = {
        "OT": "X", "BO": "X", "RE": "X", "IN": "X", "IS": "X", "NO": "X"
    }

    def gerar_folha_oficinas(self, visitas: List[Dict],
                             arquivo: Optional[str] = None) -> str:
        """
        Gera Folha de Acompanhamento de Oficinas (Frente e Verso)
        Baseado no modelo FUNCABES - Feedback de Oficina
        Dados preenchidos a partir dos registros de visita
        """
        if arquivo is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            arquivo = f"folha_oficinas_{timestamp}.docx"

        caminho_completo = os.path.join(self.pasta_relatorios, arquivo)

        doc = Document()

        # Configurar margens menores
        for section in doc.sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1)
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)

        for idx, visita in enumerate(visitas):
            if idx > 0:
                doc.add_page_break()

            # ============ FRENTE - FEEDBACK DE OFICINA ============
            self._gerar_frente_folha(doc, visita)

            # Quebra de pagina para o verso
            doc.add_page_break()

            # ============ VERSO - GRADE DE AVALIACAO ============
            self._gerar_verso_folha(doc, visita)

        doc.save(caminho_completo)
        return caminho_completo

    def _gerar_cabecalho_funcabes(self, doc):
        """Gera cabecalho institucional FUNCABES com logos"""
        # Tenta usar logos se existirem
        logo_funcabes = os.path.join('static', 'img', 'logo_funcabes.png')
        logo_cidade = os.path.join('static', 'img', 'logo_taubate.png')

        header_table = doc.add_table(rows=1, cols=3)
        header_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Logo esquerda (FUNCABES)
        cell_logo_l = header_table.cell(0, 0)
        cell_logo_l.width = Cm(3)
        if os.path.exists(logo_funcabes):
            p_logo = cell_logo_l.paragraphs[0]
            p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_logo = p_logo.add_run()
            run_logo.add_picture(logo_funcabes, width=Cm(2.5))
        else:
            self._add_cell_text(cell_logo_l, "FUNCABES", bold=True, size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)

        # Texto institucional central
        cell_centro = header_table.cell(0, 1)
        cell_centro.width = Cm(12)
        p_centro = cell_centro.paragraphs[0]
        p_centro.alignment = WD_ALIGN_PARAGRAPH.CENTER

        texto_inst = (
            "Instituida pela resolucao No12/82 do Conselho Universitario da UNITAU\n"
            "CNPJ:51.637.593/0001-32\n"
            "Av. Nove de Julho, 245, Centro, Taubate/SP- CEP:12020-200\n"
            "TEL:3633-3855 - www.funcabes.com.br - funcabes@uol.com.br"
        )
        run_centro = p_centro.add_run(texto_inst)
        run_centro.font.size = Pt(7)
        run_centro.font.name = 'Arial'

        # Logo direita (Cidade)
        cell_logo_r = header_table.cell(0, 2)
        cell_logo_r.width = Cm(3)
        if os.path.exists(logo_cidade):
            p_logo_r = cell_logo_r.paragraphs[0]
            p_logo_r.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_logo_r = p_logo_r.add_run()
            run_logo_r.add_picture(logo_cidade, width=Cm(2.5))
        else:
            self._add_cell_text(cell_logo_r, "TAUBATE", bold=True, size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    def _gerar_frente_folha(self, doc, visita):
        """Gera a frente da folha (Feedback de Oficina)"""

        # === CABECALHO INSTITUCIONAL ===
        self._gerar_cabecalho_funcabes(doc)

        doc.add_paragraph()

        # === TITULO ===
        titulo_table = doc.add_table(rows=1, cols=1)
        titulo_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell_titulo = titulo_table.cell(0, 0)
        self._set_cell_shading(cell_titulo, "2E75B6")
        self._set_cell_border(cell_titulo)

        p_titulo = cell_titulo.paragraphs[0]
        p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_titulo = p_titulo.add_run("FEEDBACK DE OFICINA")
        run_titulo.font.size = Pt(13)
        run_titulo.font.bold = True
        run_titulo.font.color.rgb = RGBColor(255, 255, 255)
        run_titulo.font.name = 'Arial'

        doc.add_paragraph()

        # === DADOS DA VISITA (preenchidos com dados reais) ===
        escola_nome_oficial = visita.get('escola_nome_oficial', visita.get('escola_nome', ''))
        data_visita = self._formatar_data(visita.get('data', ''))
        turno = visita.get('turno', '')
        mediador_nome = visita.get('mediador_nome', '')
        articulador_nome = visita.get('articulador_nome', '')
        gestor_nome = visita.get('gestor_nome', '')

        dados_table = doc.add_table(rows=6, cols=2)
        dados_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        campos = [
            ("Unidade Escolar:", escola_nome_oficial),
            ("Gestor(a):", gestor_nome),
            ("Articulador(a):", articulador_nome),
            ("Profissional:", mediador_nome),
            ("Data:", data_visita),
            ("Turno:", turno),
        ]

        for i, (label, valor) in enumerate(campos):
            cell_label = dados_table.cell(i, 0)
            cell_valor = dados_table.cell(i, 1)

            self._set_cell_border(cell_label)
            self._set_cell_border(cell_valor)
            self._set_cell_shading(cell_label, "D9E2F3")

            self._add_cell_text(cell_label, label, bold=True, size=10)
            self._add_cell_text(cell_valor, valor, size=10)

            cell_label.width = Cm(4)
            cell_valor.width = Cm(14)

        doc.add_paragraph()

        # === AREAS DE TEXTO (preenchidas com dados da visita) ===
        areas = [
            ("OBSERVACOES DAS OFICINAS", visita.get('observacoes', '')),
            ("CONTRIBUICOES / SUGESTOES", visita.get('contribuicoes', '')),
            ("COMBINADOS", visita.get('combinados', '')),
        ]

        for titulo_area, conteudo in areas:
            # Titulo da area
            area_header = doc.add_table(rows=1, cols=1)
            area_header.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell_h = area_header.cell(0, 0)
            self._set_cell_shading(cell_h, "2E75B6")
            self._set_cell_border(cell_h)
            p_h = cell_h.paragraphs[0]
            p_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_h = p_h.add_run(titulo_area)
            run_h.font.size = Pt(10)
            run_h.font.bold = True
            run_h.font.color.rgb = RGBColor(255, 255, 255)
            run_h.font.name = 'Arial'

            # Area de conteudo
            area_body = doc.add_table(rows=1, cols=1)
            area_body.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell_b = area_body.cell(0, 0)
            self._set_cell_border(cell_b)

            if conteudo:
                self._add_cell_text(cell_b, conteudo, size=10)
            else:
                for _ in range(3):
                    p_line = cell_b.add_paragraph()
                    p_line.paragraph_format.space_before = Pt(0)
                    p_line.paragraph_format.space_after = Pt(0)

            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(2)
            spacer.paragraph_format.space_after = Pt(2)

        # === ASSINATURAS ===
        doc.add_paragraph()
        assin_table = doc.add_table(rows=3, cols=3)
        assin_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Linhas de assinatura
        for col in range(3):
            cell_linha = assin_table.cell(0, col)
            p_linha = cell_linha.paragraphs[0]
            p_linha.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_linha = p_linha.add_run("_" * 25)
            run_linha.font.size = Pt(9)
            run_linha.font.name = 'Arial'

        # Nomes
        nomes_assin = [mediador_nome, articulador_nome, gestor_nome]
        for col, nome in enumerate(nomes_assin):
            cell_nome = assin_table.cell(1, col)
            p_nome = cell_nome.paragraphs[0]
            p_nome.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_nome = p_nome.add_run(nome)
            run_nome.font.size = Pt(8)
            run_nome.font.name = 'Arial'

        # Labels
        labels_assin = ["Funcionario", "Articulador(a)", "Gestor(a)"]
        for col, label in enumerate(labels_assin):
            cell_label = assin_table.cell(2, col)
            p_label = cell_label.paragraphs[0]
            p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_label = p_label.add_run(label)
            run_label.font.size = Pt(9)
            run_label.font.bold = True
            run_label.font.name = 'Arial'

    def _gerar_verso_folha(self, doc, visita):
        """Gera o verso da folha (Grade de Avaliacao) com dados preenchidos"""

        escola_nome_oficial = visita.get('escola_nome_oficial', visita.get('escola_nome', ''))
        articulador_nome = visita.get('articulador_nome', '')
        gestor_nome = visita.get('gestor_nome', '')
        data_visita = self._formatar_data(visita.get('data', ''))
        oficina = visita.get('oficina', '')
        turmas = visita.get('turmas', [])

        # === CABECALHO ===
        self._gerar_cabecalho_funcabes(doc)

        doc.add_paragraph()

        # === DADOS DO VERSO ===
        info_table = doc.add_table(rows=2, cols=2)
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        campos_info = [
            ("Articulador(a):", articulador_nome),
            ("Data:", data_visita),
            ("Unidade Escolar:", escola_nome_oficial),
            ("Oficina:", oficina),
        ]

        for i in range(2):
            for j in range(2):
                idx = i * 2 + j
                cell = info_table.cell(i, j)
                self._set_cell_border(cell)
                label, valor = campos_info[idx]
                self._add_cell_text(cell, f"{label} {valor}", bold=False, size=9)

        doc.add_paragraph()

        # === LEGENDA DA ESCALA ===
        legenda_table = doc.add_table(rows=1, cols=6)
        legenda_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        escalas = [
            ("OTIMO", "6"),
            ("BOM", "5 a 4"),
            ("REG", "3 a 2"),
            ("INSAT", "1"),
            ("INSUF", "0"),
            ("NAO Obs", "X"),
        ]

        for col, (nome, valor) in enumerate(escalas):
            cell_esc = legenda_table.cell(0, col)
            self._set_cell_border(cell_esc)
            self._set_cell_shading(cell_esc, "D9E2F3")
            p_esc = cell_esc.paragraphs[0]
            p_esc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_esc = p_esc.add_run(f"{nome}\n({valor})")
            run_esc.font.size = Pt(7)
            run_esc.font.bold = True
            run_esc.font.name = 'Arial'

        doc.add_paragraph()

        # Garante pelo menos 2 slots de turma
        while len(turmas) < 2:
            turmas.append({'turma': '', 'num_estudantes': '', 'tema': '', 'avaliacao': {}})

        # === INFO DAS TURMAS ===
        turma_info_table = doc.add_table(rows=3, cols=2)
        turma_info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for t_idx in range(2):
            turma = turmas[t_idx] if t_idx < len(turmas) else {}
            cell_header = turma_info_table.cell(0, t_idx)
            self._set_cell_border(cell_header)
            self._set_cell_shading(cell_header, "2E75B6")
            p_h = cell_header.paragraphs[0]
            p_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_h = p_h.add_run(f"TURMA {t_idx + 1}")
            run_h.font.size = Pt(9)
            run_h.font.bold = True
            run_h.font.color.rgb = RGBColor(255, 255, 255)
            run_h.font.name = 'Arial'

            cell_turma = turma_info_table.cell(1, t_idx)
            self._set_cell_border(cell_turma)
            turma_nome = turma.get('turma', '')
            num_est = turma.get('num_estudantes', '')
            self._add_cell_text(cell_turma, f"Turma: {turma_nome}    No Estudantes: {num_est}", size=8)

            cell_tema = turma_info_table.cell(2, t_idx)
            self._set_cell_border(cell_tema)
            tema = turma.get('tema', '')
            self._add_cell_text(cell_tema, f"Tema: {tema}", size=8)

        doc.add_paragraph()

        # === GRADE DE AVALIACAO COM DADOS ===
        escala_cols = ["OT", "BO", "RE", "IN", "IS", "NO"]

        # 14 colunas: Foco(1) + Indicador(1) + Turma1(6) + Turma2(6)
        grade_table = doc.add_table(rows=0, cols=14)
        grade_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Linha 1: FOCO | INDICADORES | TURMA 1 (merge 6) | TURMA 2 (merge 6)
        row1 = grade_table.add_row()
        cell_foco_h = row1.cells[0]
        self._set_cell_shading(cell_foco_h, "2E75B6")
        self._set_cell_border(cell_foco_h)
        self._add_cell_text(cell_foco_h, "FOCO", bold=True, size=7, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        cell_foco_h.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

        cell_ind_h = row1.cells[1]
        self._set_cell_shading(cell_ind_h, "2E75B6")
        self._set_cell_border(cell_ind_h)
        self._add_cell_text(cell_ind_h, "INDICADORES", bold=True, size=7, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        cell_ind_h.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

        # TURMA 1 - merge cells 2-7
        for c in range(2, 8):
            self._set_cell_border(row1.cells[c])
            self._set_cell_shading(row1.cells[c], "2E75B6")
        cell_t1 = row1.cells[2]
        cell_t1.merge(row1.cells[7])
        t1_label = turmas[0].get('turma', 'TURMA 1') if turmas[0].get('turma') else 'TURMA 1'
        self._add_cell_text(cell_t1, t1_label, bold=True, size=7, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        cell_t1.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

        # TURMA 2 - merge cells 8-13
        for c in range(8, 14):
            self._set_cell_border(row1.cells[c])
            self._set_cell_shading(row1.cells[c], "2E75B6")
        cell_t2 = row1.cells[8]
        cell_t2.merge(row1.cells[13])
        t2_label = turmas[1].get('turma', 'TURMA 2') if turmas[1].get('turma') else 'TURMA 2'
        self._add_cell_text(cell_t2, t2_label, bold=True, size=7, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        cell_t2.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

        # Linha 2: sub-headers OT/BO/RE/IN/IS/NO
        row2 = grade_table.add_row()
        self._set_cell_border(row2.cells[0])
        self._set_cell_shading(row2.cells[0], "D9E2F3")
        self._set_cell_border(row2.cells[1])
        self._set_cell_shading(row2.cells[1], "D9E2F3")

        for turma_offset in [2, 8]:
            for i, sh in enumerate(escala_cols):
                cell_sh = row2.cells[turma_offset + i]
                self._set_cell_border(cell_sh)
                self._set_cell_shading(cell_sh, "D9E2F3")
                self._add_cell_text(cell_sh, sh, bold=True, size=6, alignment=WD_ALIGN_PARAGRAPH.CENTER)

        # Dados dos indicadores com avaliacoes preenchidas
        for foco_nome, inds in self.INDICADORES_AVALIACAO:
            first_row_idx = len(grade_table.rows)

            for ind_key, indicador_texto in inds:
                row_data = grade_table.add_row()

                cell_foco = row_data.cells[0]
                self._set_cell_border(cell_foco)

                cell_ind = row_data.cells[1]
                self._set_cell_border(cell_ind)
                self._add_cell_text(cell_ind, indicador_texto, size=7)

                # Preenche avaliacao para cada turma
                for t_idx, turma_offset in enumerate([2, 8]):
                    turma_aval = {}
                    if t_idx < len(turmas):
                        turma_aval = turmas[t_idx].get('avaliacao', {})

                    valor_selecionado = turma_aval.get(ind_key, '')

                    for col_i, escala_val in enumerate(escala_cols):
                        cell_val = row_data.cells[turma_offset + col_i]
                        self._set_cell_border(cell_val)
                        if valor_selecionado == escala_val:
                            self._add_cell_text(cell_val, "X", bold=True, size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                            self._set_cell_shading(cell_val, "C6EFCE")

            # Merge das celulas do FOCO
            if len(inds) > 1:
                first_cell = grade_table.rows[first_row_idx].cells[0]
                last_cell = grade_table.rows[first_row_idx + len(inds) - 1].cells[0]
                first_cell.merge(last_cell)
            else:
                first_cell = grade_table.rows[first_row_idx].cells[0]

            self._set_cell_shading(first_cell, "E2EFDA")
            self._add_cell_text(first_cell, foco_nome, bold=True, size=7)

        # === ASSINATURAS ===
        doc.add_paragraph()
        doc.add_paragraph()
        assin_table = doc.add_table(rows=3, cols=2)
        assin_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for col in range(2):
            cell_linha = assin_table.cell(0, col)
            p_linha = cell_linha.paragraphs[0]
            p_linha.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_linha = p_linha.add_run("_" * 30)
            run_linha.font.size = Pt(9)
            run_linha.font.name = 'Arial'

        nomes_assin = [articulador_nome, gestor_nome]
        for col, nome in enumerate(nomes_assin):
            cell_nome = assin_table.cell(1, col)
            p_nome = cell_nome.paragraphs[0]
            p_nome.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_nome = p_nome.add_run(nome)
            run_nome.font.size = Pt(8)
            run_nome.font.name = 'Arial'

        labels_assin = ["Articulador(a)", "Gestor(a)"]
        for col, label in enumerate(labels_assin):
            cell_label = assin_table.cell(2, col)
            p_label = cell_label.paragraphs[0]
            p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_label = p_label.add_run(label)
            run_label.font.size = Pt(9)
            run_label.font.bold = True
            run_label.font.name = 'Arial'
